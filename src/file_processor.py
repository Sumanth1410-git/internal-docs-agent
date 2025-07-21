import os
import sys
import requests
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import logging
from datetime import datetime
import mimetypes
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.document_loaders import UnstructuredWordDocumentLoader

# Add src to path for relative imports
current_dir = Path(__file__).parent
sys.path.insert(0, str(current_dir))

# Import with error handling
try:
    from document_loader import OptimizedDocumentLoader
    DOCUMENT_LOADER_AVAILABLE = True
except ImportError:
    try:
        # Try different import path
        from .document_loader import OptimizedDocumentLoader
        DOCUMENT_LOADER_AVAILABLE = True
    except ImportError:
        print("⚠️ Document loader not available, using basic processing")
        DOCUMENT_LOADER_AVAILABLE = False

from langchain.schema import Document

logger = logging.getLogger(__name__)

class SlackFileProcessor:
    """Advanced file processing system for Slack uploads with fallback processing"""
    
    def __init__(self, bot_token: str):
        self.bot_token = bot_token
        self.temp_dir = Path("temp_uploads")
        self.temp_dir.mkdir(exist_ok=True)
        
        # Supported file types and processors
        self.supported_types = {
            '.txt': self._process_text_file,
            '.pdf': self._process_pdf_file,
            '.docx': self._process_word_file,
            '.doc': self._process_word_file,
            '.md': self._process_markdown_file,
            '.csv': self._process_csv_file,
            '.json': self._process_json_file
        }
        
        # File size limits (in MB)
        self.max_file_size = 10  # 10MB limit
        
        logger.info(f"✅ File processor initialized with {len(self.supported_types)} supported file types")
        
    def download_slack_file(self, file_info: Dict) -> Optional[Path]:
        """Download file from Slack with authentication"""
        try:
            file_url = file_info['url_private_download']
            file_name = file_info['name']
            file_size = file_info.get('size', 0)
            
            # Check file size
            if file_size > (self.max_file_size * 1024 * 1024):
                logger.warning(f"File {file_name} too large: {file_size} bytes")
                return None
            
            # Download with authentication
            headers = {'Authorization': f'Bearer {self.bot_token}'}
            response = requests.get(file_url, headers=headers, stream=True)
            
            if response.status_code == 200:
                # Create safe filename
                safe_filename = self._create_safe_filename(file_name)
                temp_file = self.temp_dir / safe_filename
                
                # Write file to disk
                with open(temp_file, 'wb') as f:
                    for chunk in response.iter_content(chunk_size=8192):
                        f.write(chunk)
                
                logger.info(f"Downloaded file: {file_name} -> {temp_file}")
                return temp_file
            else:
                logger.error(f"Failed to download file: HTTP {response.status_code}")
                return None
                
        except Exception as e:
            logger.error(f"File download failed: {e}")
            return None
    
    def process_uploaded_file(self, file_info: Dict, user_id: str) -> Tuple[bool, str, List[Document]]:
        """Process uploaded file and extract documents"""
        try:
            file_name = file_info['name']
            file_ext = Path(file_name).suffix.lower()
            
            # Check if file type is supported
            if file_ext not in self.supported_types:
                return False, f"❌ File type `{file_ext}` not supported. Supported types: {', '.join(self.supported_types.keys())}", []
            
            # Download file (for real implementation - for testing, create mock file)
            if file_info.get('url_private_download'):
                temp_file = self.download_slack_file(file_info)
                if not temp_file:
                    return False, "❌ Failed to download file from Slack", []
            else:
                # Create mock file for testing
                temp_file = self._create_mock_file(file_name, file_ext)
            
            # Process file based on type
            processor = self.supported_types[file_ext]
            documents = processor(temp_file, file_info, user_id)
            
            # Clean up temp file
            if temp_file.exists():
                temp_file.unlink()
            
            if documents:
                success_msg = f"✅ Successfully processed `{file_name}`\n"
                success_msg += f"📄 Extracted {len(documents)} document chunks\n"
                success_msg += f"📊 Total characters: {sum(len(doc.page_content) for doc in documents):,}\n"
                success_msg += f"🏷️ Added to knowledge base with source tracking"
                
                return True, success_msg, documents
            else:
                return False, f"⚠️ No content could be extracted from `{file_name}`", []
                
        except Exception as e:
            logger.error(f"File processing error: {e}")
            return False, f"❌ Error processing file: {str(e)}", []
    
    def _create_mock_file(self, file_name: str, file_ext: str) -> Path:
        """Create mock file for testing purposes"""
        mock_content = {
            '.txt': "This is a test document uploaded to the system.\nIt contains sample content for testing file processing capabilities.",
            '.md': "# Test Document\n\n## Overview\nThis is a **sample markdown** document for testing.\n\n- Feature 1\n- Feature 2",
            '.csv': "name,department,role\nJohn Doe,Engineering,Developer\nJane Smith,HR,Manager",
            '.json': '{"test": true, "content": "Sample JSON document", "features": ["upload", "processing"]}'
        }
        
        content = mock_content.get(file_ext, "Sample content for testing file processing.")
        
        temp_file = self.temp_dir / f"mock_{file_name}"
        with open(temp_file, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return temp_file
    
    def _process_text_file(self, file_path: Path, file_info: Dict, user_id: str) -> List[Document]:
        """Process plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Create document with metadata
            doc = Document(
                page_content=content,
                metadata={
                    'source_file': file_info['name'],
                    'file_type': '.txt',
                    'uploaded_by': user_id,
                    'upload_date': datetime.now().isoformat(),
                    'file_size': len(content),
                    'slack_file_id': file_info.get('id', 'test_file')
                }
            )
            
            # Split into chunks
            try:
                from langchain.text_splitter import RecursiveCharacterTextSplitter
                splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
                documents = splitter.split_documents([doc])
            except ImportError:
                # Fallback: simple chunking
                documents = self._simple_chunk_document(doc, 500, 50)
            
            return documents
            
        except Exception as e:
            logger.error(f"Text file processing failed: {e}")
            return []
    
    def _process_pdf_file(self, file_path: Path, file_info: Dict, user_id: str) -> List[Document]:
        """Process PDF files"""
        try:
            try:
                from langchain_community.document_loaders import PyPDFLoader
                loader = PyPDFLoader(str(file_path))
                documents = loader.load()
            except ImportError:
                # Fallback: try pypdf directly
                try:
                    import pypdf
                    with open(file_path, 'rb') as f:
                        reader = pypdf.PdfReader(f)
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text() + "\n"
                    
                    documents = [Document(
                        page_content=text,
                        metadata={'source_file': file_info['name'], 'file_type': '.pdf'}
                    )]
                except ImportError:
                    # Ultimate fallback
                    logger.warning("PDF processing libraries not available")
                    return [Document(
                        page_content=f"PDF file: {file_info['name']} (content extraction not available)",
                        metadata={'source_file': file_info['name'], 'file_type': '.pdf'}
                    )]
            
            # Add metadata to all documents
            for doc in documents:
                doc.metadata.update({
                    'source_file': file_info['name'],
                    'file_type': '.pdf',
                    'uploaded_by': user_id,
                    'upload_date': datetime.now().isoformat(),
                    'slack_file_id': file_info.get('id', 'test_file')
                })
            
            return documents
            
        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            return []
    
    def _process_word_file(self, file_path: Path, file_info: Dict, user_id: str) -> List[Document]:
        """Process Word documents"""
        try:
            try:
                from langchain_community.document_loaders import UnstructuredWordDocumentLoader
                loader = UnstructuredWordDocumentLoader(str(file_path))
                documents = loader.load()
            except ImportError:
                # Fallback: try python-docx
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(file_path)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    
                    documents = [Document(
                        page_content=text,
                        metadata={'source_file': file_info['name'], 'file_type': Path(file_info['name']).suffix.lower()}
                    )]
                except ImportError:
                    logger.warning("Word processing libraries not available")
                    return [Document(
                        page_content=f"Word document: {file_info['name']} (content extraction not available)",
                        metadata={'source_file': file_info['name'], 'file_type': Path(file_info['name']).suffix.lower()}
                    )]
            
            # Add metadata
            for doc in documents:
                doc.metadata.update({
                    'source_file': file_info['name'],
                    'file_type': Path(file_info['name']).suffix.lower(),
                    'uploaded_by': user_id,
                    'upload_date': datetime.now().isoformat(),
                    'slack_file_id': file_info.get('id', 'test_file')
                })
            
            return documents
            
        except Exception as e:
            logger.error(f"Word document processing failed: {e}")
            return []
    
    def _process_markdown_file(self, file_path: Path, file_info: Dict, user_id: str) -> List[Document]:
        """Process Markdown files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            doc = Document(
                page_content=content,
                metadata={
                    'source_file': file_info['name'],
                    'file_type': '.md',
                    'uploaded_by': user_id,
                    'upload_date': datetime.now().isoformat(),
                    'slack_file_id': file_info.get('id', 'test_file')
                }
            )
            
            # Split markdown content
            try:
                from langchain.text_splitter import MarkdownTextSplitter
                splitter = MarkdownTextSplitter(chunk_size=500, chunk_overlap=50)
                documents = splitter.split_documents([doc])
            except ImportError:
                # Fallback: simple chunking
                documents = self._simple_chunk_document(doc, 500, 50)
            
            return documents
            
        except Exception as e:
            logger.error(f"Markdown processing failed: {e}")
            return []
    
    def _process_csv_file(self, file_path: Path, file_info: Dict, user_id: str) -> List[Document]:
        """Process CSV files as structured data"""
        try:
            try:
                import pandas as pd
                # Read CSV
                df = pd.read_csv(file_path)
                
                # Convert to text representation
                content = f"CSV Data from {file_info['name']}:\n\n"
                content += f"Columns: {', '.join(df.columns)}\n"
                content += f"Rows: {len(df)}\n\n"
                content += df.to_string(max_rows=100)  # Limit for performance
            except ImportError:
                # Fallback: basic CSV processing
                import csv
                content = f"CSV Data from {file_info['name']}:\n\n"
                with open(file_path, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    lines = list(reader)
                    content += f"Rows: {len(lines)}\n\n"
                    for i, row in enumerate(lines[:10]):  # Show first 10 rows
                        content += f"Row {i+1}: {', '.join(row)}\n"
            
            doc = Document(
                page_content=content,
                metadata={
                    'source_file': file_info['name'],
                    'file_type': '.csv',
                    'uploaded_by': user_id,
                    'upload_date': datetime.now().isoformat(),
                    'slack_file_id': file_info.get('id', 'test_file')
                }
            )
            
            return [doc]
            
        except Exception as e:
            logger.error(f"CSV processing failed: {e}")
            return []
    
    def _process_json_file(self, file_path: Path, file_info: Dict, user_id: str) -> List[Document]:
        """Process JSON files"""
        try:
            import json
            
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert JSON to readable text
            content = f"JSON Data from {file_info['name']}:\n\n"
            content += json.dumps(data, indent=2, ensure_ascii=False)
            
            doc = Document(
                page_content=content,
                metadata={
                    'source_file': file_info['name'],
                    'file_type': '.json',
                    'uploaded_by': user_id,
                    'upload_date': datetime.now().isoformat(),
                    'slack_file_id': file_info.get('id', 'test_file')
                }
            )
            
            return [doc]
            
        except Exception as e:
            logger.error(f"JSON processing failed: {e}")
            return []
    
    def _simple_chunk_document(self, document: Document, chunk_size: int, chunk_overlap: int) -> List[Document]:
        """Simple document chunking fallback"""
        content = document.page_content
        chunks = []
        
        for i in range(0, len(content), chunk_size - chunk_overlap):
            chunk_content = content[i:i + chunk_size]
            if chunk_content.strip():
                chunk_doc = Document(
                    page_content=chunk_content,
                    metadata={**document.metadata, 'chunk_index': len(chunks)}
                )
                chunks.append(chunk_doc)
        
        return chunks
    
    def _create_safe_filename(self, filename: str) -> str:
        """Create safe filename for temporary storage"""
        # Remove dangerous characters
        safe_chars = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789.-_"
        safe_name = ''.join(c for c in filename if c in safe_chars)
        
        # Add timestamp to avoid conflicts
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        name_parts = safe_name.rsplit('.', 1)
        if len(name_parts) == 2:
            return f"{name_parts[0]}_{timestamp}.{name_parts[1]}"
        else:
            return f"{safe_name}_{timestamp}"
    
    def cleanup_temp_files(self, max_age_hours: int = 24):
        """Clean up old temporary files"""
        try:
            import time
            current_time = time.time()
            
            for file_path in self.temp_dir.glob("*"):
                if file_path.is_file():
                    file_age = current_time - file_path.stat().st_mtime
                    if file_age > (max_age_hours * 3600):  # Convert hours to seconds
                        file_path.unlink()
                        logger.info(f"Cleaned up old temp file: {file_path.name}")
                        
        except Exception as e:
            logger.error(f"Temp file cleanup failed: {e}")

# Test function
def test_file_processor():
    """Test the file processor functionality"""
    print("=== Testing Slack File Processor ===")
    
    try:
        # Initialize processor
        processor = SlackFileProcessor('test_token')
        print(f"✅ File processor initialized")
        print(f"📁 Supported types: {list(processor.supported_types.keys())}")
        
        # Test file processing with mock data
        mock_file_info = {
            'name': 'test_document.txt',
            'id': 'test_file_123',
            'size': 1024
        }
        
        success, message, documents = processor.process_uploaded_file(mock_file_info, 'test_user')
        
        print(f"\n📄 File Processing Test:")
        print(f"Success: {success}")
        print(f"Message: {message}")
        print(f"Documents: {len(documents)} chunks created")
        
        if documents:
            print(f"\n📋 Sample Document:")
            sample_doc = documents[0]
            print(f"Content preview: {sample_doc.page_content[:100]}...")
            print(f"Metadata: {sample_doc.metadata}")
        
        print(f"\n✅ File Processor Test Completed!")
        return True
        
    except Exception as e:
        print(f"❌ File processor test failed: {e}")
        return False

if __name__ == "__main__":
    test_file_processor()
